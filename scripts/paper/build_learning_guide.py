"""Generate a comprehensive learning-guide .docx for the hydro_tl_ews project.

Run: python scripts/build_learning_guide.py
Output: docs/hydro_tl_ews_Learning_Guide.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

DOC = Document()

# ----------------------------------------------------------------- base styles
normal = DOC.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = DOC.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def h2(text):
    p = DOC.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def h3(text):
    return DOC.add_heading(text, level=3)


def p(text="", bold=False, italic=False, size=11, color=None):
    para = DOC.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


def bullets(items):
    for it in items:
        para = DOC.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = para.add_run(it[0] + " ")
            r.bold = True
            para.add_run(it[1])
        else:
            para.add_run(it)


def numbered(items):
    for it in items:
        para = DOC.add_paragraph(style="List Number")
        if isinstance(it, tuple):
            r = para.add_run(it[0] + " ")
            r.bold = True
            para.add_run(it[1])
        else:
            para.add_run(it)


def table(headers, rows, widths=None):
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1F4E79")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    DOC.add_paragraph()
    return t


def callout(title, lines, fill="EAF1FB"):
    """A shaded single-cell 'box' used for Questions/Parameters/Key-idea panels."""
    t = DOC.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    cell.text = ""
    tp = cell.paragraphs[0]
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(10.5)
    tr.font.color.rgb = ACCENT
    for ln in lines:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.1)
        run = para.add_run(("•  " + ln) if not ln.startswith(("Q", "1", "2", "3", "4", "5", "6", "7", "8", "9")) else ln)
        run.font.size = Pt(10)
    DOC.add_paragraph()
    return t


def pagebreak():
    DOC.add_page_break()


# ====================================================================== COVER
title = DOC.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Mastering Transfer-Learning Deep Learning\nfor Hydrological Early Warning")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT

sub = DOC.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Complete, From-Scratch Course Built Around the `hydro_tl_ews` Project")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

DOC.add_paragraph()
meta = DOC.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "EA-LSTM • CAMELS-US • Regional Pre-training • Fine-tuning • Walk-forward Backtesting • "
    "Flood/Drought Early Warning • SHAP Explainability"
).font.size = Pt(11)

DOC.add_paragraph()
intro = DOC.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.add_run(
    "Read this top-to-bottom to understand the project, the science, the code, and the broader field — "
    "and to build durable expertise in ML/AI for water resources."
).italic = True

pagebreak()

# ================================================================ HOW TO USE
h1("0. How to Use This Guide")
p("This document is three things at once:")
numbered([
    ("A course.", " Modules are ordered so each builds on the previous one. If you are new, read in order."),
    ("A reference.", " Tables of parameters, metrics, file maps, and equations you can return to."),
    ("A project log.", " Exactly what every part of the hydro_tl_ews codebase does, and the real bugs we fixed."),
])
p("Each module contains four recurring panels:")
bullets([
    ("Key idea —", " the one thing you must internalize."),
    ("Questions to ask —", " if you can answer these, you understand the module. Use them to quiz yourself or an advisor."),
    ("Parameters to check —", " the knobs that matter and how to reason about them."),
    ("Map to code —", " which file/function implements the concept, so theory and practice stay connected."),
])
callout("How to study (suggested rhythm)", [
    "Week-by-week: 1–2 modules per week, ~6–10 weeks total for fluency.",
    "After each module, run the matching code with --smoke or a small config and watch the numbers.",
    "Keep a lab notebook: for every run, log the config, the metric, and one sentence of interpretation.",
    "Re-derive at least one equation per module by hand (NSE, KGE, z-score, NSE-loss weighting).",
])

pagebreak()

# ================================================================ BIG PICTURE
h1("1. The Big Picture — What This Project Does and Why")
p("Plain-language statement of the problem:", bold=True)
p(
    "Many rivers we care about (for floods, droughts, water supply) have little or no historical "
    "streamflow data. Physically-based models need lots of calibration data; a model trained on one "
    "river usually transfers poorly to another. This project asks: can a single deep-learning model, "
    "pre-trained on hundreds of well-gauged catchments, transfer its 'hydrological knowledge' to a "
    "data-scarce target basin and produce skillful streamflow forecasts and early warnings of floods "
    "and droughts?"
)
p("The strategy, in one sentence:", bold=True)
p(
    "Pre-train a rainfall–runoff neural network (an EA-LSTM) on the large CAMELS-US dataset, then adapt "
    "it to a target basin with a little local data (transfer learning), and evaluate it the way an "
    "operational forecaster would — rolling forward in time (walk-forward), issuing warnings against "
    "flood/drought thresholds, and explaining the drivers with SHAP."
)
callout("Key idea", [
    "This is a regional (many-basin) → local (one-basin) transfer-learning pipeline for rainfall–runoff "
    "modeling, evaluated as an operational early-warning system rather than as a static train/test split.",
])
h2("1.1 The end-to-end flow (mental model)")
table(
    ["Phase", "What happens", "Why it matters"],
    [
        ["1. Pre-train", "Train EA-LSTM on ~200–670 donor basins (CAMELS-US)", "Learns general rainfall–runoff behavior + how static attributes modulate it"],
        ["2. Zero-shot", "Apply the pre-trained model to the target basin, no tuning", "Baseline: how far does pure transfer get you?"],
        ["3. Fine-tune", "Adapt to the target with a short local record (conservative / progressive)", "Specializes the model to the target basin"],
        ["4. Walk-forward", "Roll through time, refit every 90 days, predict next chunk", "Mimics real operations; no future leakage"],
        ["5. Early warning", "Convert predicted flow to flood/drought warning probabilities vs thresholds", "Turns a regression into a decision tool"],
        ["6. Explain (SHAP)", "Attribute predictions to input features", "Trust, physical sanity-check, science"],
    ],
    widths=[1.1, 3.0, 2.6],
)
callout("Questions to ask", [
    "Q1. Why is transfer learning attractive specifically in hydrology (what is scarce, what is abundant)?",
    "Q2. What does 'data-scarce basin' mean concretely (how many years of record)?",
    "Q3. Why evaluate with walk-forward instead of a random train/test split?",
    "Q4. What is the difference between a streamflow forecast and an early-warning probability?",
])

pagebreak()

# =========================================================== LEARNING PATH MAP
h1("2. The Learning Path (Course Map)")
p(
    "Below is the recommended order. The left column is the module in this document; the right column "
    "names the external skills you should pick up in parallel. Don't try to learn everything at once — "
    "follow the dependency order."
)
table(
    ["#", "Module (this guide)", "Learn in parallel (external)"],
    [
        ["3", "Hydrology foundations", "Rainfall–runoff, the water balance, hydrographs, return periods"],
        ["4", "The data: CAMELS-US", "Meteorological forcings, catchment attributes, units"],
        ["5", "ML & time-series foundations", "Supervised learning, train/val/test, overfitting, normalization"],
        ["6", "Sequence models & LSTM", "RNN, vanishing gradients, gates, the LSTM cell"],
        ["7", "EA-LSTM architecture", "Entity-aware input gate, embeddings, Kratzert 2019"],
        ["8", "Preprocessing pipeline", "Windowing, z-score, look-ahead bias, NaN handling"],
        ["9", "Training mechanics", "Loss functions, Adam, gradient clipping, early stopping"],
        ["10", "Transfer learning", "Freezing, differential LR, fine-tuning strategies"],
        ["11", "Evaluation metrics", "NSE, KGE, PBIAS — what each rewards/penalizes"],
        ["12", "Walk-forward backtesting", "Rolling-origin CV, leakage, bias correction"],
        ["13", "Early-warning system", "Thresholds, ROC/AUC, Brier, reliability, lead time"],
        ["14", "Explainability (SHAP)", "Shapley values, feature attribution"],
        ["15", "Codebase tour", "The actual files and how data flows through them"],
        ["16", "Parameters & hyperparameters", "Every knob and how to set it"],
        ["17", "Debugging log (real bugs)", "Memory, eager evaluation, SHAP shapes, empty loaders"],
        ["18", "Running everything", "Configs, CLI, reproducing results"],
        ["19", "Becoming an expert", "Papers, books, courses, the wider field"],
        ["21", "Case study: the protocol incident", "Evaluation integrity, leakage forensics, honest reporting"],
        ["22", "The multi-target study (current frontier)", "Generalization claims, ablations, regime diversity"],
    ],
    widths=[0.3, 2.5, 3.6],
)

pagebreak()

# ================================================================ MODULE 3
h1("3. Module — Hydrology Foundations")
p(
    "You cannot evaluate a streamflow model without understanding what streamflow is and what drives it. "
    "This module is the physical vocabulary the rest of the project assumes."
)
h2("3.1 Core concepts")
bullets([
    ("Catchment / basin / watershed —", " the land area that drains to a single outlet (the gauge). Everything is 'per basin'."),
    ("Streamflow / discharge (Q) —", " volume of water passing the gauge per unit time. Measured in cubic feet/sec (cfs) or m³/s; here converted to mm/day (depth over the catchment area) so basins of different sizes are comparable."),
    ("Forcings —", " the meteorological inputs that drive runoff: precipitation, temperature, radiation, etc."),
    ("Water balance —", " precipitation = evapotranspiration + runoff + storage change. The model implicitly learns this."),
    ("Hydrograph —", " the time series of Q; has a baseflow (slow) and quickflow (storm response) component."),
    ("Snowmelt regime —", " in snowy basins, flow peaks come from melting snow (driven by temperature, radiation, day length), not just rain. The target basin here (Merced River 11264500) is snowmelt-dominated — which is why SHAP found day length & radiation dominant."),
    ("Return period / frequency —", " a 'q95' flow is exceeded 5% of the time; floods are defined by high percentiles (q95, q99), droughts by low (q5)."),
])
h2("3.2 Why static catchment attributes matter")
p(
    "Two catchments with identical rainfall can produce very different flow because of their physical "
    "character: steep vs flat, forested vs bare, sandy vs clay soils, snowy vs arid climate. The model "
    "is given 27 static attributes per basin so it can condition its rainfall–runoff mapping on basin "
    "identity — this is the heart of 'entity-aware' modeling."
)
callout("Key idea", [
    "Streamflow = f(recent weather sequence, catchment physical character). The dynamic forcings provide "
    "the weather; the static attributes provide the character.",
])
callout("Questions to ask", [
    "Q1. Why convert discharge from cfs to mm/day? What does that normalization buy you?",
    "Q2. What physically causes a flood peak in a rain-dominated vs a snowmelt-dominated basin?",
    "Q3. Which static attributes would you expect to matter most for flashiness (fast storm response)?",
    "Q4. What is baseflow, and why might a model under-predict it (recall the zero-shot PBIAS of -71%)?",
])
callout("What to read", [
    "Addor et al. (2017), 'The CAMELS dataset: catchment attributes...' — defines the attributes you use.",
    "Any intro hydrology text chapter on the rainfall–runoff process and the hydrograph.",
])

pagebreak()

# ================================================================ MODULE 4
h1("4. Module — The Data (CAMELS-US)")
p(
    "CAMELS-US (Catchment Attributes and Meteorology for Large-sample Studies) is the dataset that makes "
    "regional deep learning possible: 671 minimally-disturbed US basins with long daily records of "
    "forcings, streamflow, and static attributes."
)
h2("4.1 The three data components")
table(
    ["Component", "What it is", "In this project"],
    [
        ["Dynamic forcings", "Daily meteorology per basin (Daymet)", "6 features (see below), the time-varying input"],
        ["Streamflow (target)", "Daily observed discharge (USGS)", "QObs converted cfs→mm/day; the thing we predict"],
        ["Static attributes", "Time-invariant basin descriptors", "27 features (topo/climate/landcover/soil/geology)"],
    ],
    widths=[1.6, 2.6, 2.8],
)
h2("4.2 The 6 dynamic forcing features")
table(
    ["Variable", "Meaning", "Hydrological role"],
    [
        ["prcp(mm/day)", "Precipitation", "Primary water input"],
        ["tmax(C)", "Max air temperature", "Snowmelt, evapotranspiration"],
        ["tmin(C)", "Min air temperature", "Freeze/thaw, ET"],
        ["srad(W/m2)", "Shortwave radiation", "Snowmelt + ET energy"],
        ["vp(Pa)", "Vapor pressure", "Humidity / ET demand"],
        ["dayl(s)", "Day length", "Seasonality, snowmelt timing"],
    ],
    widths=[1.4, 2.2, 3.0],
)
h2("4.3 The 27 static attributes (5 groups)")
bullets([
    ("Topography (3):", " elev_mean, slope_mean, area_gages2"),
    ("Climate (9):", " p_mean, pet_mean, p_seasonality, frac_snow, aridity, high_prec_freq, high_prec_dur, low_prec_freq, low_prec_dur"),
    ("Land cover (5):", " frac_forest, lai_max, lai_diff, gvf_max, gvf_diff"),
    ("Soil (8):", " soil_depth_pelletier, soil_depth_statsgo, soil_porosity, soil_conductivity, max_water_content, sand_frac, silt_frac, clay_frac"),
    ("Geology (2):", " carbonate_rocks_frac, geol_permeability"),
])
h2("4.4 Data layout on disk (what the loader expects)")
p("camels_root = data/ ; the loader walks these subfolders:")
bullets([
    "data/basin_dataset_public_v1p2/basin_mean_forcing/daymet/<HUC>/<id>_lump_cida_forcing_leap.txt",
    "data/basin_dataset_public_v1p2/usgs_streamflow/<HUC>/<id>_streamflow_qc.txt",
    "data/camels_attributes_v2.0/camels_{topo,clim,hydro,vege,soil,geol}.txt  (semicolon-separated)",
])
callout("Parameters to check", [
    "camels_root — must point at the folder that directly contains basin_dataset_public_v1p2/.",
    "Date coverage — CAMELS series here END 2014-12-31. Any config period beyond 2014 yields EMPTY data (a bug we hit).",
    "Basin id format — zero-padded 8-digit USGS gauge ids (e.g. '11264500'). '08313000' is NOT in CAMELS.",
])
callout("Questions to ask", [
    "Q1. Why use a 'large-sample' dataset for pre-training rather than one long single-basin record?",
    "Q2. What is the unit of the target after conversion, and why mm/day?",
    "Q3. Which attributes are static vs dynamic, and why must statics be handled differently in the model?",
    "Q4. How would missing days or negative precipitation values corrupt training if not cleaned?",
])

pagebreak()

# ================================================================ MODULE 5
h1("5. Module — Machine-Learning & Time-Series Foundations")
p("The minimum ML you need before the LSTM makes sense.")
bullets([
    ("Supervised learning —", " learn a function f(X)→y from examples. Here X = a 365-day forcing window + static attributes, y = streamflow on the last day."),
    ("Features vs target —", " inputs vs the thing predicted. Keep them strictly separated in time to avoid leakage."),
    ("Train / validation / test —", " fit on train, tune/early-stop on validation, report on untouched test. In hydrology, splits must respect time."),
    ("Overfitting / underfitting —", " memorizing noise vs failing to learn signal. Watch the train-vs-val gap (we saw mild overfitting after epoch 23: train kept dropping, val plateaued)."),
    ("Normalization —", " rescale inputs so no feature dominates by magnitude; essential for gradient-based training. Fit scaling on TRAIN only."),
    ("Regularization —", " dropout, weight decay, early stopping — techniques that fight overfitting."),
    ("Batching & epochs —", " an epoch is one pass over the data; batches are subsets processed at once (batch_size=256 here)."),
])
h2("5.1 Time-series specifics")
bullets([
    ("Autocorrelation —", " today's flow is highly correlated with yesterday's; random shuffling of days leaks information."),
    ("Windowing / sliding window —", " convert a long series into many (sequence → target) samples. Sequence length here = 365 days."),
    ("Look-ahead bias —", " using any information from the future (including normalization statistics) invalidates results."),
    ("Stationarity & regime —", " climate/landcover can drift; one reason to refit over time (walk-forward)."),
])
callout("Key idea", [
    "Hydrological ML is time-series ML with hard rules: never let the future leak into the past — not in "
    "the split, not in the normalizer, not in the windowing.",
])
callout("Questions to ask", [
    "Q1. Why is a random train/test split dangerous for autocorrelated streamflow?",
    "Q2. Where could look-ahead bias sneak in besides the data split? (hint: the Normalizer)",
    "Q3. Why fit the normalizer on the training period only?",
    "Q4. How do batch size and learning rate interact?",
])

pagebreak()

# ================================================================ MODULE 6
h1("6. Module — Sequence Models and the LSTM")
p(
    "Streamflow on a given day depends on weeks-to-months of antecedent weather (soil moisture, snowpack "
    "have memory). A model that remembers is needed: the Recurrent Neural Network, and specifically the LSTM."
)
h2("6.1 From RNN to LSTM")
bullets([
    ("RNN —", " processes a sequence step by step, carrying a hidden state. Struggles with long memory (vanishing gradients)."),
    ("LSTM (Long Short-Term Memory) —", " adds a cell state and gates (input, forget, output) that learn what to remember, forget, and expose. Solves long-range memory."),
    ("Why LSTM for hydrology —", " catchment storage (snow, soil, groundwater) behaves like the LSTM cell state: it integrates inputs over time and releases slowly."),
])
h2("6.2 The gates (intuition)")
table(
    ["Gate", "Question it answers", "Hydrological analogy"],
    [
        ["Forget", "What to drop from memory?", "Soil drying out between storms"],
        ["Input", "What new info to store?", "New rainfall entering storage"],
        ["Cell state", "The running memory", "Total catchment storage"],
        ["Output", "What to expose now?", "How much storage becomes flow"],
    ],
    widths=[1.2, 2.6, 2.8],
)
p(
    "Note on the forget gate bias: the project sets initial_forget_bias = 3.0. Initializing the forget "
    "gate 'open' makes the LSTM remember by default early in training — a standard trick that helps it "
    "learn long dependencies (Kratzert 2019)."
)
callout("Questions to ask", [
    "Q1. What problem do the LSTM gates solve that a vanilla RNN cannot?",
    "Q2. Why is the cell state a good analogy for catchment storage?",
    "Q3. What does a high initial forget-gate bias do to early training, and why help here?",
    "Q4. Why can't an LSTM be parallelized across time (relevant to the ~18% GPU utilization we saw)?",
])
callout("What to read", [
    "Hochreiter & Schmidhuber (1997) — the original LSTM (skim).",
    "Christopher Olah, 'Understanding LSTM Networks' (blog) — the best visual intuition.",
])

pagebreak()

# ================================================================ MODULE 7
h1("7. Module — The EA-LSTM Architecture (the heart of the model)")
p(
    "A plain LSTM treats every basin the same. The Entity-Aware LSTM (EA-LSTM, Kratzert et al. 2019) lets "
    "the static catchment attributes control how the network reads the weather — one shared model that "
    "behaves differently per basin."
)
h2("7.1 The key modification")
p(
    "In a standard LSTM, the input gate depends on the dynamic input. In an EA-LSTM, the input gate is "
    "computed ONCE per basin from the static attributes only, and held fixed across all time steps. The "
    "dynamic forcings flow through the rest of the cell. Intuition: the static attributes set a per-basin "
    "'filter' that decides which aspects of the weather matter for that catchment."
)
table(
    ["Piece", "Depends on", "Role"],
    [
        ["Static input gate i", "Static attributes (27)", "Per-basin, time-constant 'character' filter"],
        ["Forget/cell/output gates", "Dynamic forcings + hidden state", "Process the weather sequence over time"],
        ["Embedding of statics", "Static attributes", "Learned compression of basin character"],
        ["Dense head", "Final hidden state", "Maps LSTM memory → streamflow value"],
    ],
    widths=[1.9, 2.3, 2.6],
)
callout("Key idea", [
    "EA-LSTM = one model, many basins. Static attributes parameterize a per-basin input gate, so the "
    "network can generalize across catchments AND specialize to each — exactly what regional transfer needs.",
])
h2("7.2 Why this enables transfer learning")
p(
    "Because basin character enters only through the static gate, a model pre-trained on hundreds of "
    "basins has already learned the general weather→flow dynamics. For a new basin you mainly need to get "
    "its static gate / head right — which is why even zero-shot transfer works partially, and a little "
    "fine-tuning works well."
)
callout("Questions to ask", [
    "Q1. In an EA-LSTM, what exactly is computed from the static attributes, and how often?",
    "Q2. Why does making the input gate static (not dynamic) aid cross-basin generalization?",
    "Q3. Which parts would you freeze vs fine-tune when adapting to a new basin? (links to Module 10)",
    "Q4. Find the EA-LSTM forward pass in src/hydro_tl_ews/models/ealstm.py — trace where statics enter.",
])
callout("Map to code", [
    "src/hydro_tl_ews/models/ealstm.py — EALSTM, EALSTMConfig, freeze_lstm(), unfreeze_lstm(), "
    "trainable_parameter_groups(). Read this alongside the paper.",
])
callout("What to read", [
    "Kratzert et al. (2019), 'Towards learning universal, regional, and local hydrological behaviors via "
    "machine learning applied to large-sample datasets', HESS — THE paper for this project.",
])

pagebreak()

# ================================================================ MODULE 8
h1("8. Module — The Preprocessing Pipeline")
p("How raw CAMELS text becomes model-ready tensors. This is where most real-world bugs live.")
h2("8.1 The steps (in order)")
numbered([
    ("Load —", " read per-basin forcing + streamflow text files; convert cfs→mm/day using basin area."),
    ("Quality control —", " drop non-physical values (negative precip/flow), linearly interpolate short gaps (≤3 days)."),
    ("Align —", " intersect forcing and streamflow dates (their indices can differ)."),
    ("Slice to period —", " restrict to the configured train/val/eval window."),
    ("Normalize —", " z-score the dynamic forcings (fit on train period only); min-max the static attributes."),
    ("Window —", " build (365-day sequence → last-day target) samples with a sliding window."),
    ("Mask —", " drop windows containing any NaN forcing or a NaN target."),
])
h2("8.2 The normalizers")
table(
    ["Normalizer", "Method", "Fit on", "Why"],
    [
        ["Dynamic (forcings)", "z-score (x-μ)/σ", "Pre-train period only", "Comparable scales; no look-ahead"],
        ["Static (attributes)", "min-max → [0,1]", "Across donor basins", "Bounded inputs to the static gate"],
    ],
    widths=[1.8, 1.7, 1.7, 2.0],
)
h2("8.3 The lazy windowing (a core engineering lesson)")
p(
    "The naïve approach materializes every 365-day window for every basin into one giant array. For full "
    "CAMELS that is ~60–70 GB and crashed the machine. The fix: store each basin's normalized forcings "
    "ONCE and cut each window on demand in __getitem__. Memory dropped ~365× (to ~1 GB) with identical "
    "samples. Validity (no-NaN window + non-NaN target) is computed with a cumulative-sum trick in O(T) "
    "memory instead of building a 3-D boolean array."
)
callout("Key idea", [
    "A Dataset should usually be lazy: keep raw arrays small, generate samples on access. Materializing "
    "all windows upfront trades a 365× memory blow-up for nothing.",
])
callout("Parameters to check", [
    "sequence_length (365) — the look-back window; must be ≤ the available record or you get ZERO samples.",
    "max_gap_days (3) — longest gap linearly interpolated; longer gaps stay NaN and drop the window.",
    "DYNAMIC_FEATURES / STATIC_ATTRIBUTES — the exact column lists the model expects.",
])
callout("Questions to ask", [
    "Q1. Where could look-ahead bias enter this pipeline, and how is it prevented?",
    "Q2. Why must a 3-month warmup with sequence_length=365 produce zero training samples?",
    "Q3. What is the memory cost of materializing all windows vs storing per-basin arrays?",
    "Q4. Why z-score for dynamics but min-max for statics?",
])
callout("Map to code", [
    "src/hydro_tl_ews/data/preprocessing.py — Normalizer, StaticNormalizer, quality_control, "
    "align_forcing_streamflow, make_sequences.",
    "src/hydro_tl_ews/data/datasets.py — MultiBasinSequenceDataset (the lazy windowing).",
    "src/hydro_tl_ews/data/camels.py — CamelsDataset loader + feature lists.",
])

pagebreak()

# ================================================================ MODULE 9
h1("9. Module — Training Mechanics")
p("How the model actually learns, and the specific choices this project makes.")
h2("9.1 The loss functions")
bullets([
    ("NSE loss (pre-training) —", " a differentiable form of (1 − Nash–Sutcliffe Efficiency), weighted by per-basin variance so high-flow basins don't dominate the gradient (Kratzert 2019). This is why each basin's std is carried through the dataset."),
    ("MSE (fine-tuning / baseline) —", " plain mean squared error for single-basin adaptation."),
])
h2("9.2 The optimizer and stabilizers")
table(
    ["Knob", "Value here", "What it does"],
    [
        ["Optimizer", "Adam", "Adaptive per-parameter learning rates"],
        ["learning_rate (head)", "1e-3", "Step size for the dense head"],
        ["lstm_lr", "1e-5", "Smaller step for pre-trained LSTM (differential LR)"],
        ["weight_decay", "0.0", "L2 regularization (off by default)"],
        ["clip_grad_norm", "1.0", "Caps gradient norm to prevent exploding updates"],
        ["dropout", "0.4", "Randomly zeroes units to fight overfitting"],
        ["batch_size", "256", "Samples per gradient step"],
    ],
    widths=[1.8, 1.2, 3.2],
)
h2("9.3 Early stopping")
p(
    "Training watches validation loss; if it does not improve for `patience` epochs, it stops and keeps "
    "the best checkpoint. In our 200-donor run, best validation was epoch 23 and later epochs only "
    "overfit (train ↓, val flat) — a textbook early-stopping signal. Note: when the validation window is "
    "shorter than sequence_length there is no val set (val=nan) and the loop falls back to train loss."
)
callout("Questions to ask", [
    "Q1. Why weight the NSE loss by per-basin variance? What goes wrong without it?",
    "Q2. Why use a smaller learning rate for the LSTM than the head during fine-tuning?",
    "Q3. What does gradient clipping protect against in a recurrent model?",
    "Q4. How do you read a train-vs-validation curve to decide when to stop?",
])
callout("Map to code", [
    "src/hydro_tl_ews/training/trainer.py — Trainer (modes, _step, run_epoch, fit, early stopping).",
    "src/hydro_tl_ews/models/losses.py — NSELoss.",
])

pagebreak()

# ================================================================ MODULE 10
h1("10. Module — Transfer Learning Strategies")
p("The project's reason for being: move knowledge from many basins to one data-scarce basin.")
table(
    ["Mode", "What is trainable", "Loss", "When to use"],
    [
        ["pretrain", "All parameters", "NSE", "Build the regional base model"],
        ["zero_shot", "Nothing (inference)", "—", "Baseline transfer with no local data"],
        ["conservative", "Dense head only (LSTM frozen)", "MSE", "Very little local data; avoid overfitting"],
        ["progressive", "Head + last 25% of LSTM, differential LR", "MSE", "A bit more local data; deeper adaptation"],
        ["local baseline", "All params, random init", "MSE", "Control: no transfer, train from scratch"],
    ],
    widths=[1.3, 2.6, 0.8, 2.3],
)
h2("10.1 Why a spectrum of fine-tuning?")
p(
    "Freezing more of the network means fewer parameters to fit, which is safer when local data is tiny "
    "(less overfitting) but less flexible. Progressive unfreezing with a small LSTM learning rate adapts "
    "the dynamics gently without destroying pre-trained knowledge (catastrophic forgetting)."
)
callout("Key idea", [
    "Match adaptation depth to data quantity: frozen head for scraps of data, progressive unfreezing for "
    "more. The local-from-scratch baseline tells you how much transfer actually bought you.",
])
h2("10.2 What the results showed (this project, corrected protocol 2026-07-02)")
p(
    "These are the Merced results under the strict data-scarce protocol (refits confined to the 2-year "
    "warmup + observed data, validated on a held-out 90-day tail — see Module 21 for why this matters)."
)
table(
    ["Setting", "NSE", "Reading"],
    [
        ["Local from-scratch (2 yr)", "-1.40", "Collapses: worse than predicting the mean. No transfer = no skill"],
        ["Zero-shot", "0.08", "Pure transfer is weak on this snowmelt basin (PBIAS -71%, big under-prediction)"],
        ["Static fine-tune A (head only)", "0.40", "Two years of local data + frozen LSTM = moderate skill"],
        ["Static fine-tune B (progressive)", "0.51", "Partial LSTM unfreezing adapts dynamics further"],
        ["Walk-forward A (refits)", "0.53", "Operational loop adds bias correction + periodic refits"],
        ["Walk-forward B (refits)", "0.57", "Best variant: progressive refits inside the loop"],
    ],
    widths=[2.4, 0.9, 3.4],
)
callout("A number to distrust", [
    "Earlier runs reported walk-forward NSE 0.76/0.82. Those runs let refits train on the full 1990+ "
    "record — 19 extra years the 'data-scarce' story says don't exist — and used no validation during "
    "refits. The drop to 0.53/0.57 under the honest protocol is the measured price of that leakage. "
    "Module 21 dissects the whole incident.",
])
callout("Questions to ask", [
    "Q1. Why does freezing the LSTM reduce overfitting on a tiny local record?",
    "Q2. What is catastrophic forgetting and how does a small lstm_lr mitigate it?",
    "Q3. What does the local-baseline control isolate?",
    "Q4. Why is the ladder (-1.40 → 0.08 → 0.40 → 0.53) the headline evidence? What does each rung isolate?",
])
callout("Map to code", [
    "src/hydro_tl_ews/training/transfer.py — fine_tune_conservative, fine_tune_progressive, "
    "train_local_baseline, FineTuneConfig.",
])

pagebreak()

# ================================================================ MODULE 11
h1("11. Module — Evaluation Metrics")
p("Three continuous metrics dominate hydrology. Know exactly what each rewards.")
table(
    ["Metric", "Range / ideal", "What it measures", "Watch-out"],
    [
        ["NSE", "(-inf, 1]; 1 best", "Skill vs predicting the mean flow", "NSE=0 means 'no better than the mean'"],
        ["KGE", "(-inf, 1]; 1 best", "Balances correlation, bias, variability", "More diagnostic than NSE"],
        ["PBIAS %", "0 best", "Systematic over/under-prediction", "Sign matters: -71% = severe under-prediction"],
    ],
    widths=[0.9, 1.5, 2.7, 1.7],
)
p(
    "NSE relates to mean-squared error normalized by the variance of observations — a model that beats "
    "the climatological mean scores >0. KGE decomposes performance into correlation, bias ratio, and "
    "variability ratio, so a low KGE tells you WHICH part is failing. PBIAS is the percent volume error."
)
callout("Questions to ask", [
    "Q1. Why can NSE be very negative, and what does negative NSE imply about the model?",
    "Q2. KGE vs NSE — when would they disagree, and why is KGE often preferred?",
    "Q3. The zero-shot model had NSE 0.08 but PBIAS -71%. Reconcile those two numbers physically.",
    "Q4. Which metric would you optimize for a water-supply forecast vs a flood-peak forecast?",
])
callout("Map to code", [
    "src/hydro_tl_ews/evaluation/metrics.py — nse, kge, pbias, auc_roc, brier_score, f1_at_threshold, "
    "reliability_curve.",
])

pagebreak()

# ================================================================ MODULE 12
h1("12. Module — Walk-Forward Backtesting")
p(
    "The honest way to evaluate an operational forecasting system. Instead of one fixed split, you roll "
    "through time: train up to a date, predict the next chunk, advance, periodically refit."
)
h2("12.1 The loop (as implemented, corrected protocol)")
numbered([
    ("Start at initial_train_end —", " fit a normalizer on data up to that date (never later)."),
    ("Refit —", " fine-tune on data from refit_train_start (the warmup start) up to 'now', on a deep copy "
     "so the base weights stay clean. The most recent val_tail_days (90) are HELD OUT of training and "
     "used as the early-stopping signal; the best-validation epoch's weights are restored at the end."),
    ("Predict the next chunk —", " refit_every_days = 90, so it forecasts ~3 months ahead."),
    ("Online bias correction —", " add a correction estimated from the previous chunk's observed-minus-"
     "predicted error (causal: never uses the chunk being predicted)."),
    ("Advance —", " move the origin forward and repeat until eval_end."),
])
p(
    "In our corrected run: initial_train_end 2010-12-31, eval 2011–2014, 17 refits, 1461 daily "
    "predictions, final NSE 0.53 (Approach A) / 0.57 (Approach B). Each refit is a small fast fine-tune "
    "(2–3 epochs). full_period reaches back to 1990 ONLY so the extreme thresholds have a >=20-year "
    "record; refit_train_start stops the training window from also reaching back there."
)
callout("Key idea", [
    "Walk-forward = rolling-origin cross-validation for time series. It forbids future leakage and "
    "measures performance the way the model would actually be deployed. But the loop is only as honest "
    "as its window boundaries — see Module 21 for how two subtle boundary mistakes inflated our own "
    "results by 0.23 NSE.",
])
callout("Parameters to check", [
    "initial_train_end — where evaluation begins (must leave enough history for one sequence window).",
    "refit_every_days (90) — how often you re-tune; smaller = more adaptive but more compute.",
    "val_tail_days (90) — held-out tail of each refit window; the early-stopping signal. 0 disables "
    "(falls back to train-loss stopping — dangerous, see Module 21).",
    "refit_train_start (2009-01-01) — earliest date refits may train on; omitting it lets refits see the "
    "whole loaded record and silently break a data-scarce simulation.",
    "online_bias_correction (true) — corrects systematic drift between refits.",
    "full_period — for thresholds you need >=20 years, hence 1990–2014.",
])
callout("Questions to ask", [
    "Q1. Why deep-copy the model before each refit?",
    "Q2. What leakage would a naïve random split allow that walk-forward forbids?",
    "Q3. How does online bias correction differ from re-training, and when does it help most?",
    "Q4. Trade-off: what happens to skill and to runtime as refit_every_days shrinks?",
    "Q5. With val_tail_days = 90 = refit_every_days, the model never trains on the quarter just before "
    "each forecast. What seasonal risk does that create, and what alternatives exist?",
])
callout("Map to code", [
    "src/hydro_tl_ews/training/walk_forward.py — walk_forward(), WalkForwardConfig, _build_loader, bias logic.",
    "scripts/stages/walk_forward_stage.py — wires data, thresholds, metrics, SHAP, outputs.",
])

pagebreak()

# ================================================================ MODULE 13
h1("13. Module — The Early-Warning System")
p(
    "A streamflow number is not a warning. This module converts predicted flow into the probability that "
    "a flood or drought threshold will be crossed, at several lead times, and scores those warnings."
)
h2("13.1 From flow to warning")
bullets([
    ("Regional thresholds (RFA) —", " compute flood/drought levels from the long record: q95 & q99 (floods), q5 (drought). Needs ≥20 years for stable estimates."),
    ("Warning labels —", " did observed flow cross the threshold within the next 1/3/7 days? (the truth)"),
    ("Warning probabilities —", " how confidently did the model's prediction cross it? (the forecast)"),
    ("Lead time —", " 1, 3, 7 days ahead — warnings are more useful but harder further out."),
])
h2("13.2 Scoring the warnings (classification metrics)")
table(
    ["Metric", "Ideal", "Meaning"],
    [
        ["AUC-ROC", "1.0", "Ranking skill: can it separate events from non-events?"],
        ["Brier score", "0.0", "Mean squared error of the probability forecast (calibration+sharpness)"],
        ["F1 @ 0.5", "1.0", "Balance of precision/recall at a 0.5 cutoff (NaN if no positives predicted)"],
        ["Reliability curve", "diagonal", "Are predicted probabilities calibrated to observed frequencies?"],
    ],
    widths=[1.4, 0.8, 4.0],
)
p(
    "Project results: flood AUC ≈ 0.98 across lead times (excellent ranking skill), low Brier; drought "
    "AUC ≈ 0.79–0.84 (harder). F1 was NaN in some cells when the model predicted no positives above 0.5 "
    "— a reminder that threshold choice matters."
)
callout("Questions to ask", [
    "Q1. Why is AUC robust to class imbalance while raw accuracy is not (floods are rare)?",
    "Q2. What does a Brier score actually penalize that AUC ignores?",
    "Q3. Why might F1 be NaN, and what does that reveal about the operating threshold?",
    "Q4. Why do longer lead times usually degrade warning skill?",
])
callout("Map to code", [
    "src/hydro_tl_ews/evaluation/extreme_thresholds.py — regional_thresholds, warning_labels, "
    "predicted_warning_probabilities.",
    "src/hydro_tl_ews/evaluation/metrics.py — auc_roc, brier_score, f1_at_threshold, reliability_curve.",
])

pagebreak()

# ================================================================ MODULE 14
h1("14. Module — Explainability (SHAP)")
p(
    "A skillful black box is not enough for science or trust. SHAP attributes each prediction to the "
    "input features, so you can check the model learned physically sensible drivers."
)
bullets([
    ("Shapley values —", " a game-theory fair-credit allocation: each feature's average marginal contribution to the prediction."),
    ("GradientExplainer —", " a SHAP method suited to differentiable models like neural nets."),
    ("Sequence reduction —", " the (365×6) forcing attributions are averaged over time to one value per feature, comparable to the static attributes."),
    ("Global importance —", " mean absolute SHAP per feature, ranked."),
])
p(
    "Project result: day length and shortwave radiation dominated — the energy drivers of snowmelt for "
    "this Sierra Nevada basin. Static attributes showed ~0 because the analysis is single-basin (they "
    "don't vary), so they cannot move a within-basin prediction. That is the physically correct outcome."
)
callout("Questions to ask", [
    "Q1. What does a Shapley value represent, and why is it considered 'fair'?",
    "Q2. Why are static-attribute SHAP values ~0 in a single-basin explanation?",
    "Q3. Does feature importance imply causation? What are the caveats?",
    "Q4. Why does day length dominating make physical sense for a snowmelt basin?",
])
callout("Map to code", [
    "src/hydro_tl_ews/xai/shap_analysis.py — explain_predictions, global_importance (and the output-axis "
    "squeeze fix we added for newer SHAP).",
])

pagebreak()

# ================================================================ MODULE 15
h1("15. Module — The Codebase Tour (file by file)")
p("Now connect every concept to the file that implements it. Read in this order.")
table(
    ["File", "Responsibility"],
    [
        ["src/hydro_tl_ews/data/camels.py", "Load CAMELS forcings/streamflow/attributes; feature lists; cfs→mm/day"],
        ["src/hydro_tl_ews/data/preprocessing.py", "Normalizers, quality control, alignment, make_sequences"],
        ["src/hydro_tl_ews/data/datasets.py", "MultiBasinSequenceDataset (lazy windowing) → PyTorch samples"],
        ["src/hydro_tl_ews/data/clustering.py", "Select similar donor basins by static-attribute similarity"],
        ["src/hydro_tl_ews/data/synthetic_camels.py", "Fake CAMELS-like data for the smoke test (no download)"],
        ["src/hydro_tl_ews/models/ealstm.py", "EA-LSTM architecture; freeze/unfreeze; param groups"],
        ["src/hydro_tl_ews/models/losses.py", "NSELoss (variance-weighted)"],
        ["src/hydro_tl_ews/training/trainer.py", "Generic train/val loop, modes, early stopping, save/load"],
        ["src/hydro_tl_ews/training/transfer.py", "Conservative/progressive fine-tune, local baseline"],
        ["src/hydro_tl_ews/training/walk_forward.py", "Rolling-origin backtester + online bias correction"],
        ["src/hydro_tl_ews/evaluation/metrics.py", "NSE/KGE/PBIAS + AUC/Brier/F1/reliability"],
        ["src/hydro_tl_ews/evaluation/extreme_thresholds.py", "RFA thresholds, warning labels & probabilities"],
        ["src/hydro_tl_ews/xai/shap_analysis.py", "SHAP attribution and global importance"],
        ["src/hydro_tl_ews/utils/{config,device,logging,seed}.py", "Config loading, GPU selection, logging, reproducibility"],
        ["scripts/run_experiment.py", "CLI entry point; dispatches to a stage by config"],
        ["scripts/stages/*.py", "One file per pipeline phase (pretrain, zero_shot, finetune, walk_forward, ...)"],
        ["scripts/smoke_pipeline.py", "Tiny end-to-end run on synthetic data"],
        ["configs/*.yaml", "All run settings; the CLI reads exactly one per run"],
    ],
    widths=[2.9, 3.5],
)
h2("15.1 How a single run flows")
numbered([
    "You run: python scripts/run_experiment.py --config configs/<stage>.yaml",
    "run_experiment.py reads the YAML, sets the seed, and dispatches by cfg.stage.",
    "The stage file loads CAMELS, builds Normalizers, constructs MultiBasinSequenceDataset → DataLoader.",
    "It builds an EALSTM, wraps it in a Trainer (or transfer/walk_forward routine), and trains/evaluates.",
    "Outputs (checkpoints, metrics JSON, CSV/parquet, SHAP) are written under results/.",
])
callout("Questions to ask", [
    "Q1. Trace one forcing window from a CAMELS .txt file to a tensor entering the model.",
    "Q2. Which file decides cpu vs cuda, and how?",
    "Q3. Where is the seed set, and what does reproducibility require besides a seed?",
    "Q4. How does the CLI know which stage to run?",
])

pagebreak()

# ================================================================ MODULE 16
h1("16. Module — Parameters & Hyperparameters Reference")
p("Every important knob, grouped. This is your tuning cheat-sheet.")
h2("16.1 Data parameters")
table(
    ["Parameter", "Typical", "Effect / how to reason"],
    [
        ["camels_root", "data", "Path to dataset root"],
        ["target_basin", "11264500", "Held out of pretrain; evaluated for transfer"],
        ["similar_donor_count", "200 / null", "How many donor basins; more = better base but slower"],
        ["exclusion_buffer_km", "50", "Drop donors near the target to avoid spatial leakage"],
        ["pretrain_period", "1980–2010", "Training years; fit normalizer here only"],
        ["validation_period", "2010–2014", "Early-stopping window"],
        ["sequence_length", "365", "Look-back days; must be <= record length"],
        ["num_workers", "0–4", "DataLoader parallelism; 0 is safest on Windows"],
    ],
    widths=[1.9, 1.2, 3.3],
)
h2("16.2 Model parameters")
table(
    ["Parameter", "Typical", "Effect"],
    [
        ["hidden_size", "256", "LSTM capacity; bigger = more powerful + slower/overfit risk"],
        ["dropout", "0.4", "Regularization strength"],
        ["initial_forget_bias", "3.0", "Encourages long memory early in training"],
    ],
    widths=[1.9, 1.2, 3.3],
)
h2("16.3 Training parameters")
table(
    ["Parameter", "Typical", "Effect"],
    [
        ["batch_size", "256", "Throughput vs gradient noise"],
        ["epochs", "30–50", "Upper bound; early stopping usually ends sooner"],
        ["patience", "10", "Epochs of no val-improvement before stopping"],
        ["learning_rate (head)", "1e-3", "Head step size"],
        ["lstm_lr", "1e-5", "Pre-trained LSTM step size (fine-tune)"],
        ["weight_decay", "0.0", "L2 regularization"],
        ["clip_grad_norm", "1.0", "Exploding-gradient guard"],
        ["unfreeze_fraction", "0.25", "Progressive: share of LSTM unfrozen"],
    ],
    widths=[1.9, 1.2, 3.3],
)
h2("16.4 Walk-forward parameters")
table(
    ["Parameter", "Typical", "Effect"],
    [
        ["initial_train_end", "2010-12-31", "Evaluation start"],
        ["eval_end", "2014-12-31", "Evaluation end (<= data end!)"],
        ["refit_every_days", "90", "Refit cadence"],
        ["val_tail_days", "90", "Held-out tail per refit = early-stopping signal (0 = train-loss stopping)"],
        ["refit_train_start", "2009-01-01", "Earliest refit training date; guards the data-scarce window"],
        ["online_bias_correction", "true", "Running drift correction"],
        ["lead_times", "[1,3,7]", "Warning horizons (days)"],
    ],
    widths=[1.9, 1.2, 3.3],
)
callout("Questions to ask (tuning intuition)", [
    "Q1. If you have only 1 year of local data, which knobs do you change first?",
    "Q2. Validation NSE plateaus while train keeps improving — which knobs address it?",
    "Q3. GPU is underused and epochs are slow — is the bottleneck data loading or the LSTM? How to tell?",
    "Q4. You need results faster on the same hardware — which parameters trade skill for speed?",
])

pagebreak()

# ================================================================ MODULE 17
h1("17. Module — Debugging Log (the real bugs we fixed)")
p(
    "Real expertise is built by fixing real failures. This session's debugging is a compact course in "
    "the kinds of bugs that bite ML/hydrology pipelines."
)
table(
    ["Symptom", "Root cause", "Fix", "Transferable lesson"],
    [
        ["MemoryError on a 22 MiB array (with 41 GB free!)", "Dataset materialized ALL windows (~60–70 GB); the small array was the last straw", "Lazy windowing in MultiBasinSequenceDataset", "Watch peak/cumulative memory, not the failing line; prefer lazy datasets"],
        ["KeyError: 'Target basin 08313000 ...'", "Target id not in CAMELS; clustering path needs a real id", "Use 11264500; understand pretrain tolerated it as a no-op", "Validate ids/keys against the actual data early"],
        ["Empty datasets / silent failures", "Config periods 2015–2020 but CAMELS ends 2014", "Remap to 2009–2014 windows", "Always check data coverage vs config dates"],
        ["KeyError on dict.get default", "get('a', d['b']) eagerly evaluates d['b']", "Use get('a') or d['b']", "Default args are always evaluated in Python"],
        ["SHAP concatenate dim mismatch", "Newer SHAP returns extra output axis (N,F,1)", "Squeeze the singleton axis", "Pin/verify library output shapes across versions"],
        ["DataLoader num_samples=0", "Warmup < sequence_length → empty dataset", "Record NaN and skip", "Guard against empty data in loops"],
        ["val=nan at every walk-forward refit (spotted by an external reviewer)", "Refits never received a validation loader → train-loss early stopping only", "val_tail_days holdout + best-epoch weight restoration in Trainer.fit", "Validation must SELECT the model, not just stop it; log files are evidence — grep them"],
        ["Walk-forward NSE 0.76 too good to be true", "Refit window sliced [:fit_end] from data start → trained on 1990+, not the 2-yr warmup", "refit_train_start config; honest rerun dropped NSE to 0.53", "Audit window BOUNDARIES, not just the split direction; inflated results feel great and are worthless"],
    ],
    widths=[1.7, 1.9, 1.5, 1.9],
)
callout("Key idea", [
    "Most pipeline bugs are not in the math — they are in data coverage, shapes, units, keys, memory, and "
    "library version drift. Build the habit of checking those FIRST.",
])
callout("Questions to ask", [
    "Q1. How would you proactively detect a 'config date beyond data range' bug before it crashes?",
    "Q2. Why did 41 GB free still produce a MemoryError? What were you really out of?",
    "Q3. What is the general Python gotcha behind the dict.get bug?",
    "Q4. How do you make a pipeline robust to library upgrades like SHAP's API change?",
])

pagebreak()

# ================================================================ MODULE 18
h1("18. Module — Running Everything (Reproduce the Results)")
h2("18.1 Environment")
bullets([
    "Create venv: python -m venv .venv ; activate it.",
    "Install: pip install torch --index-url https://download.pytorch.org/whl/cu124 ; then numpy pandas pyyaml scikit-learn shap matplotlib pyarrow ; pip install -e . (and pytest for tests).",
    "Verify GPU: python -c \"import torch; print(torch.cuda.is_available())\".",
])
h2("18.2 Smoke test first (no CAMELS needed)")
bullets([
    "python scripts/run_experiment.py --config configs/smoke_test.yaml --smoke",
    "pytest tests/ -q   (expect: 14 passed)",
])
h2("18.3 Full pipeline order")
numbered([
    "pretrain_subset200.yaml  (or the heavier pretrain.yaml)  → pretrain checkpoint",
    "zero_shot.yaml  → zero_shot_metrics.json",
    "finetune_conservative.yaml  → finetune_conservative.pt",
    "finetune_progressive.yaml  → finetune_progressive.pt",
    "local_baseline.yaml  → local_baseline.pt",
    "walk_forward.yaml  → walk_forward.parquet, metrics, warnings, SHAP",
    "min_data_sensitivity.yaml  → min_data_sensitivity.csv",
])
p(
    "Multi-target study (once results/checkpoints/pretrain.pt exists): "
    "python scripts/run_multi_target.py — runs all six stages for each of the seven target basins and "
    "writes results/multi_target/summary.csv. Resumable: completed stages are skipped. "
    "Configs are generated by scripts/gen_multi_target_configs.py (Module 22)."
)
callout("Parameters to check before a long run", [
    "Does every config date fall within the basin's record (<= 2014)?",
    "Do downstream configs point at a checkpoint that actually exists?",
    "Is sequence_length <= the shortest warmup window you expect to use?",
    "Is the GPU visible, and is num_workers safe for your OS?",
])

pagebreak()

# ================================================================ MODULE 19
h1("19. Becoming an Expert — The Wider Field & Reading Path")
h2("19.1 Foundational papers (read in this order)")
numbered([
    "Olah (2015), 'Understanding LSTM Networks' (blog) — intuition for gates.",
    "Kratzert et al. (2018), 'Rainfall–runoff modelling using LSTM networks', HESS — LSTM for streamflow.",
    "Kratzert et al. (2019), EA-LSTM, HESS — the architecture in this project.",
    "Addor et al. (2017), CAMELS attributes, HESS — the data.",
    "Newman et al. (2015), CAMELS forcings/flow — the data.",
    "Gupta et al. (2009), 'Decomposition of the mean squared error...' — the KGE metric.",
    "Nash & Sutcliffe (1970) — the NSE metric (classic).",
    "Lundberg & Lee (2017), 'A Unified Approach to Interpreting Model Predictions' (SHAP).",
    "Nearing et al. (2021), 'What role does hydrological science play in the age of ML?' — perspective.",
    "Kratzert et al. (2021/2024), large-sample & ungauged-basin LSTM follow-ups — the frontier.",
])
h2("19.2 Books & courses")
bullets([
    ("Deep learning —", " Goodfellow et al., 'Deep Learning'; or fast.ai / Andrew Ng courses for momentum."),
    ("Time series —", " Hyndman & Athanasopoulos, 'Forecasting: Principles and Practice' (free online)."),
    ("Hydrology —", " any standard engineering hydrology text (Chow, Maidment & Mays) for the physics."),
    ("ML for Earth science —", " the emerging 'Deep Learning for the Earth Sciences' literature; NeuralHydrology docs/tutorials."),
])
h2("19.3 Tools to master")
bullets([
    "PyTorch (models, autograd, DataLoaders) — you used all three here.",
    "pandas / numpy (time-series wrangling) — the backbone of preprocessing.",
    "NeuralHydrology (the community library that productionizes exactly this kind of pipeline).",
    "xarray + ERA5/Daymet (gridded meteorology) for going beyond CAMELS.",
    "Experiment tracking (Weights & Biases / MLflow) and configuration discipline (Hydra/YAML).",
])
h2("19.4 Skills checklist — 'I have command when I can...'")
bullets([
    "Explain EA-LSTM and why it enables regional → local transfer, without notes.",
    "Derive NSE/KGE and predict how a change in bias/variance moves them.",
    "Design a leakage-free, walk-forward evaluation for a new basin.",
    "Diagnose whether a slow run is data-bound or compute-bound, and fix it.",
    "Take a new region (your own gauge), assemble forcings + attributes, and run this pipeline end-to-end.",
    "Defend the physical plausibility of a model using SHAP and domain knowledge.",
])
callout("Capstone project (do this to cement mastery)", [
    "Pick a real basin you care about (CAMELS or your own USGS gauge).",
    "Assemble its forcings + static attributes; choose donors; pretrain or reuse a checkpoint.",
    "Run zero-shot → fine-tune → walk-forward; produce flood/drought warnings with AUC/Brier.",
    "Explain drivers with SHAP and write a 2-page report defending every modeling choice.",
])

pagebreak()

# ================================================================ GLOSSARY
h1("20. Glossary (quick reference)")
table(
    ["Term", "One-line meaning"],
    [
        ["EA-LSTM", "LSTM whose input gate is set per-basin from static attributes"],
        ["CAMELS-US", "671-basin US dataset: forcings, streamflow, attributes"],
        ["Forcings", "Time-varying meteorological inputs (precip, temp, radiation, ...)"],
        ["Static attributes", "Time-invariant basin descriptors (soil, topography, climate...)"],
        ["NSE", "Nash–Sutcliffe Efficiency; skill vs predicting the mean"],
        ["KGE", "Kling–Gupta Efficiency; correlation+bias+variability"],
        ["PBIAS", "Percent bias; systematic over/under-prediction"],
        ["Zero-shot", "Applying a model with no target-specific training"],
        ["Fine-tuning", "Adapting a pre-trained model to a new task/basin"],
        ["Walk-forward", "Rolling-origin time-series backtesting"],
        ["Online bias correction", "Running adjustment of predictions from recent errors"],
        ["RFA thresholds", "Flow percentiles (q5/q95/q99) defining drought/flood"],
        ["Lead time", "How many days ahead a warning is issued"],
        ["AUC-ROC", "Ranking skill of a probabilistic classifier"],
        ["Brier score", "MSE of probabilistic forecasts (lower better)"],
        ["SHAP", "Shapley-value feature attribution for predictions"],
        ["Look-ahead bias", "Leaking future information into training/normalization"],
        ["Sequence length", "Number of past days fed to the LSTM (here 365)"],
        ["Donor basin", "A source basin used for regional pre-training"],
        ["Exclusion buffer", "Radius around a target within which donors are removed (here 50 km)"],
        ["Refit window", "The date range a walk-forward refit is allowed to train on"],
        ["Best-weight restoration", "Reloading the best-validation-epoch weights after early stopping"],
        ["BSS", "Brier Skill Score: improvement of a forecast's Brier over a reference (climatology)"],
        ["DOY climatology", "Day-of-year event frequency from history; the benchmark a model must beat"],
    ],
    widths=[1.8, 4.6],
)

pagebreak()

# ================================================================ MODULE 21
h1("21. Case Study — The Protocol Incident (a course in evaluation integrity)")
p(
    "On 2026-07-02 an external reader's one-paragraph caveat triggered a forensic audit that changed the "
    "project's headline number from NSE 0.76 to NSE 0.53. Nothing in the model changed — only the "
    "honesty of the evaluation. Study this module carefully: it is the most transferable lesson in the "
    "whole project, and reviewers of your paper WILL probe exactly these seams."
)
h2("21.1 The caveat, verbatim (paraphrased)")
p(
    "“During walk-forward evaluation the code retrains the model but passes no validation data into "
    "the training pipeline. The loop falls back to the training loss to decide when to stop. Training "
    "loss is always decreasing, so the model lacks the unseen test it needs to know when to stop — it "
    "risks memorizing the exact weather of that short window instead of learning general patterns.”"
)
h2("21.2 The forensic method (how the claim was verified)")
numbered([
    ("Read the code path —", " walk_forward() built loaders; refit_fn(model, train_loader, val_loader=None, ...) "
     "— so Trainer.fit fell back to train-loss stopping. Claim structurally TRUE."),
    ("Check the artifacts, not just the code —", " grep 'val=' in the run logs: every refit line said "
     "val=nan in BOTH published runs. Claim empirically TRUE for every published number."),
    ("Quantify the blast radius —", " which results depended on the flawed loop? Both walk-forward "
     "variants. Static fine-tunes, zero-shot and local baseline were clean."),
    ("Hunt for adjacent flaws —", " auditing the same function found a SECOND, worse issue nobody had "
     "flagged: the refit training slice was [:fit_end] from the START of the loaded record (1990), not "
     "from the warmup. The 'data-scarce' refits were training on 21–25 years of local streamflow."),
])
h2("21.3 Diagnosis: two distinct flaws, opposite directions")
table(
    ["Flaw", "Mechanism", "Direction of damage"],
    [
        ["No validation in refits", "Train-loss early stopping cannot detect overfitting; no best-epoch selection", "Risk of memorization (small here: head-only refits train 257 params for 3 epochs)"],
        ["Refit window reached to 1990", "Refits saw two decades of 'nonexistent' local data", "Inflated skill: the headline was not a data-scarce result at all"],
    ],
    widths=[1.6, 2.9, 2.4],
)
p(
    "Note the irony: the reviewer's proposed MECHANISM (memorizing a short window) was wrong — the "
    "window wasn't short, which is precisely why the results looked so good. But the reviewer's INSTINCT "
    "(this loop is not evaluating what you think it is) was exactly right. Take vague unease seriously; "
    "verify the specific mechanism yourself."
)
h2("21.4 The fixes (all regression-tested)")
bullets([
    ("val_tail_days (90) —", " each refit holds out its most recent 90 days as a genuine validation set."),
    ("Best-weight restoration —", " Trainer.fit snapshots the best-validation epoch and restores it; "
     "validation now selects the model instead of merely stopping it."),
    ("refit_train_start (2009-01-01) —", " refits may only train on warmup + observed evaluation data."),
    ("Archived, not deleted —", " the old artifacts live in results/archive_2026-06_pre_protocol_fix/ so "
     "the before/after comparison is reproducible evidence, not an anecdote."),
])
h2("21.5 What the honest rerun showed")
table(
    ["Quantity", "Old (inflated)", "Corrected", "Reading"],
    [
        ["Walk-forward A: NSE / KGE", "0.763 / 0.804", "0.527 / 0.625", "0.23 NSE = measured value of 19 yr of local data"],
        ["Walk-forward B: NSE / KGE", "0.819 / 0.804", "0.571 / 0.584", "B stays the best variant"],
        ["Flood EWS AUC", "0.98–0.99", "0.95–0.98", "Survives; still beats DOY climatology (BSS +0.31..+0.80)"],
        ["Peak underestimation", "~28%", "~46%", "The flattering number was borrowed data, now disclosed"],
        ["Bias-correction effect", "0.72 → 0.76", "0.33 → 0.53", "Correction is now load-bearing, not cosmetic"],
    ],
    widths=[2.0, 1.4, 1.4, 2.6],
)
callout("Key idea", [
    "An evaluation protocol is a claim about what information the model had. Audit the BOUNDARIES of "
    "every training window, not just the direction of the split. And when honest numbers drop, they do "
    "not weaken the paper — they become the paper: the 0.23-NSE gap directly quantifies the value of "
    "local record length, which was a stated research objective.",
])
callout("Questions to ask", [
    "Q1. Why did per-epoch validation not exist even though the Trainer supported it? Trace the None.",
    "Q2. Why was the reviewer's memorization mechanism implausible for head-only refits? (Count the parameters.)",
    "Q3. Which published claims survived the correction unchanged, and what protected them?",
    "Q4. Why is archiving the flawed artifacts better science than deleting them?",
    "Q5. Grep exercise: find the log lines proving all 17 corrected refits used a validation tail.",
])
callout("Map to code", [
    "src/hydro_tl_ews/training/walk_forward.py — val_tail_days, refit_train_start, window slicing.",
    "src/hydro_tl_ews/training/trainer.py — fit(), restore_best_weights.",
    "tests/test_training_fixes.py — the four regression tests locking all of this in.",
    "docs/DOCX_CORRECTIONS_REQUIRED.md item 10 + PROJECT_STATUS_AND_CONTINUATION.txt — the audit trail.",
])

pagebreak()

# ================================================================ MODULE 22
h1("22. The Multi-Target Study — Where the Project Is Going")
p(
    "A single basin makes a case study; a framework claim needs evidence across hydrologic regimes. "
    "While the full pretrain runs, understand the design you are about to execute."
)
h2("22.1 The three claims under construction")
numbered([
    ("Generalization —", " the same corrected protocol on 7 targets spanning 6 regimes. If transfer helps "
     "in most and the failures are physically explicable, the method is a framework, not an anecdote."),
    ("Donor-selection ablation —", " Merced evaluated with BOTH pretrains: 199 similarity-selected donors "
     "vs the full 649-donor model. Directly tests the Ougahi & Rowan hypothesis cited in the proposal."),
    ("Value of data —", " data-scarce vs full-history refit protocols differ by 0.23 NSE on Merced — an "
     "empirical answer to 'how much is a longer local record worth?' (proposal Objective 6)."),
])
h2("22.2 The seven targets (and why each is there)")
table(
    ["Basin", "Where / regime", "What it tests"],
    [
        ["11264500 Merced R, CA", "Sierra snowmelt", "The anchor; connects to all prior results"],
        ["09107000 Taylor R, CO", "Rockies snowmelt, 3300 m", "Does snowmelt skill replicate in a different mountain range?"],
        ["14222500 EF Lewis R, WA", "Pacific NW maritime rain", "Winter-rain regime; strong seasonality, no snow signal"],
        ["02128000 Little R, NC", "Humid Southeast rain", "Convective rainfall-runoff; the 'easy' regime"],
        ["01544500 Kettle Ck, PA", "Northeast mixed snow/rain", "Transitional regime; rain-on-snow events"],
        ["11224500 Los Gatos Ck, CA", "Semi-arid, flashy/ephemeral", "The deliberate hard case: expect degraded skill; the regime-boundary discussion"],
        ["05507600 Lick Ck, MO", "Continental plains, summer-peak", "Thunderstorm-driven; tests the regional model's convective knowledge"],
    ],
    widths=[1.7, 1.9, 3.3],
)
p(
    "Selection method: of 671 basins, 231 passed the completeness screen (>=99% daily flow 1990–2014, "
    "gap-free 2009–2014); one representative per regime bin was chosen, preferring ~250–400 km² areas "
    "and mutual separation >50 km. Full screen table: results/multi_target_screen.csv."
)
h2("22.3 The leakage decision that had to happen BEFORE training")
p(
    "Every future target must be excluded from the pretrain (plus a 50 km buffer), otherwise zero-shot "
    "and transfer claims are contaminated. The first pretrain launch excluded only Merced; it was killed "
    "two hours in and relaunched with all seven targets excluded (649 donors remain). Two hours of GPU "
    "time bought seven leakage-free targets — the cheapest correctness fix this project will ever make. "
    "Lesson: leakage prevention is a DESIGN-TIME decision; after training it is unfixable without a rerun."
)
h2("22.4 Execution (one command)")
bullets([
    "scripts/gen_multi_target_configs.py — regenerates the 42 per-basin stage configs (7 basins × 6 stages).",
    "scripts/run_multi_target.py — gates on results/checkpoints/pretrain.pt, runs every stage per basin, "
    "skips completed work (resumable), writes results/multi_target/summary.csv.",
    "Per basin: finetune A, finetune B, local baseline, zero-shot, walk-forward A, walk-forward B — all "
    "under the corrected protocol of Module 21.",
])
callout("Predictions to write down BEFORE the results arrive (good science hygiene)", [
    "P1. Little R (humid SE) will show the best transfer skill; rainfall-runoff dominates CAMELS donors.",
    "P2. Los Gatos Ck (ephemeral) will be the worst — possibly negative NSE — because zero-flow days and "
    "flashy events violate the regimes the regional model saw most.",
    "P3. Taylor R will roughly replicate Merced's pattern (snowmelt physics transfers).",
    "P4. The full-671 pretrain will beat subset-200 on non-snowmelt targets but only tie it on Merced "
    "(similarity selection already covered snowmelt donors).",
    "Write your own predictions now; compare when summary.csv lands. Being wrong is informative.",
])
callout("Questions to ask", [
    "Q1. Why must the exclusion buffer apply to EVERY target, not just the one being evaluated?",
    "Q2. Why is one shared pretrain for all targets acceptable (vs a per-target pretrain), and what "
    "does the answer cost in rigor?",
    "Q3. If Los Gatos fails as predicted, is that a weakness of the paper or a contribution? Argue both.",
    "Q4. What would it take to upgrade from 'case study' to 'framework' in a reviewer's eyes — and does "
    "7 basins reach it?",
])
callout("Map to code", [
    "scripts/stages/pretrain_stage.py — exclude_targets_and_buffer (multi-target leakage guard).",
    "scripts/gen_multi_target_configs.py — TARGETS dict, protocol constants, config templates.",
    "scripts/run_multi_target.py — orchestration, resumability, summary aggregation.",
])

# ------------------------------------------------------------------ save
out_dir = Path(__file__).resolve().parents[2] / "docs"
out_dir.mkdir(parents=True, exist_ok=True)
out_path = out_dir / "hydro_tl_ews_Learning_Guide.docx"
DOC.save(out_path)
print(f"Saved: {out_path}")
print(f"Sections: 22 modules; paragraphs: {len(DOC.paragraphs)}")
